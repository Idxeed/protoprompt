"""Bounded, local-only text and office document ingestion."""

from __future__ import annotations

from collections.abc import Iterator
from contextlib import contextmanager
from dataclasses import dataclass
import hashlib
import mimetypes
from pathlib import Path
import threading
from typing import Any
from zipfile import BadZipFile, ZipFile

from protoprompt.rag.types import Document


class DocumentReadError(ValueError):
    """A document was unsupported, malformed, encrypted, or over a limit."""


@dataclass(frozen=True)
class ReaderLimits:
    """Resource bounds applied before and during parsing."""

    max_bytes: int = 10 * 1024 * 1024
    max_chars: int = 5_000_000
    max_pages: int = 500
    max_archive_entries: int = 2_000
    max_archive_uncompressed_bytes: int = 50 * 1024 * 1024
    max_archive_ratio: float = 100.0
    max_pdf_stream_bytes: int = 25 * 1024 * 1024

    def __post_init__(self) -> None:
        for name, value in vars(self).items():
            if value <= 0:
                raise ValueError(f"{name} must be positive")


_TEXT_SUFFIXES = {
    ".txt", ".md", ".markdown", ".rst", ".org",
    ".py", ".pyi", ".js", ".jsx", ".ts", ".tsx", ".java", ".kt",
    ".go", ".rs", ".c", ".h", ".cc", ".cpp", ".cxx", ".hpp",
    ".cs", ".rb", ".php", ".swift", ".scala", ".sh", ".bash",
    ".zsh", ".fish", ".ps1", ".sql", ".graphql", ".yaml", ".yml",
    ".toml", ".ini", ".cfg", ".conf", ".json", ".jsonl", ".xml",
    ".csv", ".tsv", ".dockerfile",
}
_HTML_SUFFIXES = {".html", ".htm"}


# pypdf 6.16 exposes decoder ceilings as module-level knobs rather than
# per-reader options. Keep changes serialized and restore the prior values so
# one bounded local read cannot relax another caller's limits.
_PDF_DECODE_LIMIT_LOCK = threading.RLock()
_PYPDF_OUTPUT_LIMITS = (
    "MAX_DECLARED_STREAM_LENGTH",
    "ZLIB_MAX_OUTPUT_LENGTH",
    "LZW_MAX_OUTPUT_LENGTH",
    "RUN_LENGTH_MAX_OUTPUT_LENGTH",
    "JBIG2_MAX_OUTPUT_LENGTH",
    "MAX_ARRAY_BASED_STREAM_OUTPUT_LENGTH",
)


@contextmanager
def _bounded_pypdf_decode(max_output_bytes: int) -> Iterator[None]:
    """Temporarily cap pypdf's compressed-stream expansion.

    ``PdfReader`` expands a page's content stream before this reader can count
    it. pypdf documents these module-level limits as its supported guard for
    that operation; use the reader's configured stream ceiling rather than its
    much larger library default.
    """
    from pypdf import filters as pypdf_filters

    with _PDF_DECODE_LIMIT_LOCK:
        previous = {
            name: getattr(pypdf_filters, name)
            for name in _PYPDF_OUTPUT_LIMITS
        }
        try:
            for name, value in previous.items():
                configured = int(value)
                # pypdf treats zero as "unlimited" for its decoder limits;
                # a bounded reader must never inherit that opt-out.
                bounded = max_output_bytes if configured <= 0 else min(
                    configured, max_output_bytes
                )
                setattr(pypdf_filters, name, bounded)
            yield
        finally:
            for name, value in previous.items():
                setattr(pypdf_filters, name, value)


class LocalDocumentReader:
    """Read allow-listed local files into :class:`~protoprompt.rag.Document`.

    URLs are deliberately rejected. ``allowed_root`` is recommended for web
    services: resolved paths (including symlinks) must remain beneath it.
    """

    def __init__(
        self,
        *,
        limits: ReaderLimits | None = None,
        allowed_root: str | Path | None = None,
        encoding: str = "utf-8",
    ) -> None:
        self.limits = limits or ReaderLimits()
        self.allowed_root = (
            Path(allowed_root).resolve(strict=True)
            if allowed_root is not None
            else None
        )
        self.encoding = encoding

    def read(
        self,
        path: str | Path,
        *,
        doc_id: str | None = None,
        metadata: dict[str, Any] | None = None,
        password: str | None = None,
    ) -> Document:
        candidate = Path(path)
        if str(path).lower().startswith(("http://", "https://", "file://")):
            raise DocumentReadError("remote URLs and URI strings are not accepted")
        try:
            resolved = candidate.resolve(strict=True)
        except (OSError, RuntimeError) as exc:
            raise DocumentReadError(f"document does not exist: {candidate}") from exc
        if not resolved.is_file():
            raise DocumentReadError(f"document is not a regular file: {candidate}")
        if self.allowed_root is not None and not resolved.is_relative_to(self.allowed_root):
            raise DocumentReadError("document resolves outside allowed_root")
        size = resolved.stat().st_size
        if size > self.limits.max_bytes:
            raise DocumentReadError(
                f"document is {size} bytes; limit is {self.limits.max_bytes}"
            )

        suffix = resolved.suffix.lower()
        if suffix in _TEXT_SUFFIXES or resolved.name.lower() in {"dockerfile", "makefile"}:
            text, reader_name = self._read_text(resolved), "text"
        elif suffix in _HTML_SUFFIXES:
            text, reader_name = self._read_html(resolved), "html"
        elif suffix == ".pdf":
            text, reader_name = self._read_pdf(resolved, password=password), "pdf"
        elif suffix == ".docx":
            text, reader_name = self._read_docx(resolved), "docx"
        else:
            raise DocumentReadError(f"unsupported document extension: {suffix or '<none>'}")

        self._check_chars(text)
        media_type = mimetypes.guess_type(resolved.name)[0] or "text/plain"
        provenance = dict(metadata or {})
        provenance.update({
            "source_kind": "local_file",
            "source_name": resolved.name,
            "source_uri": resolved.as_uri(),
            "media_type": media_type,
            "byte_size": size,
            "reader": reader_name,
        })
        identity = doc_id or "file_" + hashlib.blake2b(
            resolved.as_uri().encode("utf-8"), digest_size=16
        ).hexdigest()
        return Document(doc_id=identity, text=text, metadata=provenance)

    def _read_text(self, path: Path) -> str:
        raw = path.read_bytes()
        if len(raw) > self.limits.max_bytes:
            raise DocumentReadError(
                f"document is {len(raw)} bytes; limit is {self.limits.max_bytes}"
            )
        if b"\x00" in raw:
            raise DocumentReadError("text/source document contains NUL bytes")
        try:
            return raw.decode(self.encoding).replace("\r\n", "\n").replace("\r", "\n")
        except UnicodeDecodeError as exc:
            raise DocumentReadError(
                f"document is not valid {self.encoding}; choose encoding explicitly"
            ) from exc

    def _read_html(self, path: Path) -> str:
        try:
            from bs4 import BeautifulSoup
        except ImportError as exc:
            raise ImportError(
                "HTML reading requires Beautiful Soup. "
                "Install with: pip install 'protoprompt[documents]'"
            ) from exc
        markup = self._read_text(path)
        soup = BeautifulSoup(markup, "html.parser")
        for tag in soup(["script", "style", "noscript", "template"]):
            tag.decompose()
        return soup.get_text("\n", strip=True)

    def _read_pdf(self, path: Path, *, password: str | None) -> str:
        try:
            from pypdf import PdfReader
            from pypdf.errors import LimitReachedError, PdfReadError
        except ImportError as exc:
            raise ImportError(
                "PDF reading requires pypdf. "
                "Install with: pip install 'protoprompt[documents]'"
            ) from exc
        try:
            # ``page.get_contents()`` performs decompression itself. Apply the
            # limit before it (and ``extract_text()``) touch a stream, rather
            # than only checking the already-expanded byte string.
            with _bounded_pypdf_decode(self.limits.max_pdf_stream_bytes):
                reader = PdfReader(path, strict=True)
                if reader.is_encrypted:
                    if password is None or not reader.decrypt(password):
                        raise DocumentReadError("PDF is encrypted; a valid password is required")
                if len(reader.pages) > self.limits.max_pages:
                    raise DocumentReadError(
                        f"PDF has {len(reader.pages)} pages; limit is {self.limits.max_pages}"
                    )
                parts: list[str] = []
                stream_total = 0
                for page in reader.pages:
                    contents = page.get_contents()
                    if contents is not None:
                        stream_total += len(contents.get_data())
                        if stream_total > self.limits.max_pdf_stream_bytes:
                            raise DocumentReadError("PDF content streams exceed the configured limit")
                    parts.append(page.extract_text() or "")
                    self._check_chars("\n\n".join(parts))
                return "\n\n".join(parts)
        except LimitReachedError as exc:
            raise DocumentReadError("PDF content streams exceed the configured limit") from exc
        except PdfReadError as exc:
            raise DocumentReadError("PDF is malformed or cannot be read") from exc

    def _read_docx(self, path: Path) -> str:
        self._check_docx_archive(path)
        try:
            from docx import Document as DocxDocument
        except ImportError as exc:
            raise ImportError(
                "DOCX reading requires python-docx. "
                "Install with: pip install 'protoprompt[documents]'"
            ) from exc
        document = DocxDocument(path)
        parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
        for table in document.tables:
            for row in table.rows:
                parts.append("\t".join(cell.text for cell in row.cells))
        return "\n".join(parts)

    def _check_docx_archive(self, path: Path) -> None:
        try:
            with ZipFile(path) as archive:
                entries = archive.infolist()
                if len(entries) > self.limits.max_archive_entries:
                    raise DocumentReadError("DOCX contains too many archive entries")
                names = {entry.filename for entry in entries}
                if "[Content_Types].xml" not in names or "word/document.xml" not in names:
                    raise DocumentReadError("file is not a valid DOCX package")
                total = sum(entry.file_size for entry in entries)
                if total > self.limits.max_archive_uncompressed_bytes:
                    raise DocumentReadError("DOCX uncompressed size exceeds the configured limit")
                compressed = max(1, sum(entry.compress_size for entry in entries))
                if total / compressed > self.limits.max_archive_ratio:
                    raise DocumentReadError("DOCX compression ratio exceeds the configured limit")
                for entry in entries:
                    if entry.filename.endswith(".rels"):
                        relationships = archive.read(entry).lower()
                        if b'targetmode="external"' in relationships:
                            raise DocumentReadError("DOCX external relationships are not allowed")
        except BadZipFile as exc:
            raise DocumentReadError("file is not a valid DOCX archive") from exc

    def _check_chars(self, text: str) -> None:
        if len(text) > self.limits.max_chars:
            raise DocumentReadError(
                f"extracted text is {len(text)} characters; limit is {self.limits.max_chars}"
            )


def read_document(
    path: str | Path,
    **kwargs: Any,
) -> Document:
    """Convenience wrapper around :class:`LocalDocumentReader`."""

    reader_keys = {"limits", "allowed_root", "encoding"}
    reader_options = {key: kwargs.pop(key) for key in list(kwargs) if key in reader_keys}
    return LocalDocumentReader(**reader_options).read(path, **kwargs)
